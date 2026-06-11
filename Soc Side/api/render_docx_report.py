from pathlib import Path
import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BASE = Path(__file__).resolve().parents[1]

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)

def set_page_margins(section, top=0.7, bottom=0.7, left=0.75, right=0.75):
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)

def set_cell_border(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'), 'single')
        elem.set(qn('w:sz'), '4')
        elem.set(qn('w:space'), '0')
        elem.set(qn('w:color'), 'B7C9E2')
        tc_borders.append(elem)
    tc_pr.append(tc_borders)

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(22, 40, 74)

def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(90, 90, 90)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 78, 121)
    run.font.size = Pt(15 if level == 1 else 12)

def add_body(doc, text, after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text if text else "N/A")
    run.font.size = Pt(10.5)

def add_bullet_list(doc, items):
    if not items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.1
        p.add_run("N/A").font.size = Pt(10.5)
        return
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.1
        run = p.add_run(str(item))
        run.font.size = Pt(10.5)

def add_key_value_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = True

    hdr = table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Value"
    set_cell_shading(hdr[0], "D9EAF7")
    set_cell_shading(hdr[1], "D9EAF7")

    for key, value in rows:
        row = table.add_row().cells
        row[0].text = str(key)
        row[1].text = str(value if value not in [None, ""] else "N/A")

    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.05
                for r in p.runs:
                    r.font.size = Pt(10)

def normalize_to_list(value):
    if value is None:
        return []
    if isinstance(value, list):
        return value
    return [str(value)]

def split_analyst_paragraphs(text):
    if not text:
        return ["N/A"]
    parts = [p.strip() for p in str(text).split("\n\n") if p.strip()]
    if parts:
        return parts

    # fallback: split long single-block report into paragraph-sized chunks
    text = str(text).strip()
    sentences = [s.strip() for s in text.replace("\n", " ").split(". ") if s.strip()]
    if len(sentences) <= 2:
        return [text]

    chunks = []
    current = []
    for s in sentences:
        current.append(s)
        if len(current) >= 2:
            chunk = ". ".join(current)
            if not chunk.endswith("."):
                chunk += "."
            chunks.append(chunk)
            current = []
    if current:
        chunk = ". ".join(current)
        if not chunk.endswith("."):
            chunk += "."
        chunks.append(chunk)
    return chunks

def add_analyst_report(doc, text):
    paragraphs = split_analyst_paragraphs(text)
    for para in paragraphs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(para)
        run.font.size = Pt(10.5)

def render_docx(run_id: str):
    src = BASE / "reports" / "json" / f"{run_id}_report.json"
    if not src.exists():
        raise FileNotFoundError(f"JSON report not found: {src}")

    data = json.loads(src.read_text(encoding="utf-8"))

    doc = Document()

    for section in doc.sections:
        set_page_margins(section)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)

    add_title(doc, "Security Incident Report")
    add_subtitle(doc, f"Run ID: {run_id}")

    add_heading(doc, "1. Executive Summary")
    add_body(doc, data.get("executive_summary", "No executive summary available."))

    add_heading(doc, "2. Plain-Language Summary")
    add_body(doc, data.get("plain_language_summary", "No summary available."))

    add_heading(doc, "3. Incident Overview")
    add_key_value_table(doc, [
        ("Verdict", data.get("verdict", "N/A")),
        ("Targeted Service", data.get("targeted_service", "N/A")),
        ("Likely Vulnerability", data.get("likely_vulnerability", "N/A")),
        ("Likely CVE", data.get("likely_cve", "N/A")),
        ("Attack Outcome", data.get("attack_outcome", "N/A")),
        ("Confidence", data.get("confidence", "N/A")),
    ])

    doc.add_paragraph()

    add_heading(doc, "4. Proven Impact")
    add_body(doc, data.get("proven_impact", "N/A"))

    add_heading(doc, "5. Why This Matters")
    add_body(doc, data.get("risk_explanation", "N/A"))

    add_heading(doc, "6. Business Impact")
    add_body(doc, data.get("business_impact", "N/A"))

    add_heading(doc, "7. Evidence Summary")
    add_bullet_list(doc, normalize_to_list(data.get("evidence_summary", [])))

    add_heading(doc, "8. Immediate Actions")
    add_bullet_list(doc, normalize_to_list(data.get("immediate_actions", [])))

    add_heading(doc, "9. Recommended Mitigations")
    add_bullet_list(doc, normalize_to_list(data.get("mitigations", [])))

    add_heading(doc, "10. Long-Term Security Improvements")
    add_bullet_list(doc, normalize_to_list(data.get("long_term_recommendations", [])))

    add_heading(doc, "11. Detailed Analyst Report")
    add_analyst_report(doc, data.get("analyst_report", "N/A"))

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run("Generated automatically by the SOC reporting pipeline.")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)

    out = BASE / "reports" / "docx" / f"{run_id}_report.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out

if __name__ == "__main__":
    output = render_docx("run_hard_004")
    print(f"saved to {output}")
